from pathlib import Path

from docx import Document


def parse_docx(path: Path) -> str:
    doc = Document(path)
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(c.text.strip() for c in row.cells))
        lines.append("\n".join(rows))
    return "\n\n".join(lines).strip()
